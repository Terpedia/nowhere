#!/usr/bin/env python3
"""Build JPS submission DOCX files from the checked-in Markdown sources."""
import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "manuscript" / "jps-formatted"
OUT.mkdir(exist_ok=True)

BLACK = "000000"
HEADER_FILL = "404040"
BORDER = "D9D9D9"
ALT_FILL = "F3F5F7"

def set_cell_shading(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = tcPr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tcPr.append(shd)
    shd.set(qn("w:fill"), fill)

def set_cell_borders(cell, color=BORDER):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    borders = tcPr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tcPr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = "w:" + edge
        node = borders.find(qn(tag))
        if node is None:
            node = OxmlElement(tag)
            borders.append(node)
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), "4")
        node.set(qn("w:space"), "0")
        node.set(qn("w:color"), color)

def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE")
    run._r.append(fld)

def configure(doc):
    sec = doc.sections[0]
    sec.top_margin = Inches(1)
    sec.bottom_margin = Inches(1)
    sec.left_margin = Inches(1)
    sec.right_margin = Inches(1)
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    normal.font.size = Pt(12)
    normal.font.color.rgb = RGBColor.from_string(BLACK)
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    normal.paragraph_format.space_after = Pt(0)
    for name, size, before, after in (("Title", 16, 0, 18), ("Heading 1", 14, 18, 6), ("Heading 2", 13, 12, 4), ("Heading 3", 12, 10, 3)):
        style = styles[name]
        style.font.name = "Times New Roman"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(BLACK)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True
        # Word's built-in Title style can carry a blue paragraph rule; remove it.
        ppr = style._element.find(qn("w:pPr"))
        if ppr is not None:
            pborder = ppr.find(qn("w:pBdr"))
            if pborder is not None:
                ppr.remove(pborder)
    add_page_number(sec.footer.paragraphs[0])

def add_inline(paragraph, text):
    # Preserve bold and italic Markdown spans without retaining Markdown syntax.
    pos = 0
    pattern = re.compile(r"(\*\*[^*]+\*\*|\*[^*]+\*|`[^`]+`)")
    for match in pattern.finditer(text):
        if match.start() > pos:
            paragraph.add_run(text[pos:match.start()])
        token = match.group(0)
        if token.startswith("**"):
            run = paragraph.add_run(token[2:-2])
            run.bold = True
        elif token.startswith("*"):
            run = paragraph.add_run(token[1:-1])
            run.italic = True
        else:
            run = paragraph.add_run(token[1:-1])
            run.font.name = "Courier New"
            run.font.size = Pt(10)
        pos = match.end()
    if pos < len(text):
        paragraph.add_run(text[pos:])

def add_table(doc, rows):
    table = doc.add_table(rows=1, cols=len(rows[0]))
    table.autofit = True
    table.style = "Table Grid"
    for i, value in enumerate(rows[0]):
        cell = table.rows[0].cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        p.paragraph_format.line_spacing = 1.0
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(value)
        r.bold = True
        r.font.name = "Times New Roman"
        r.font.size = Pt(9)
        r.font.color.rgb = RGBColor(255, 255, 255)
        set_cell_shading(cell, HEADER_FILL)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        set_cell_borders(cell)
    for row_i, row in enumerate(rows[1:], 1):
        cells = table.add_row().cells
        for col_i, value in enumerate(row):
            cell = cells[col_i]
            cell.text = ""
            p = cell.paragraphs[0]
            p.paragraph_format.line_spacing = 1.0
            p.paragraph_format.space_after = Pt(0)
            add_inline(p, value)
            for run in p.runs:
                run.font.name = "Times New Roman"
                run.font.size = Pt(8.5)
            if row_i % 2 == 0:
                set_cell_shading(cell, ALT_FILL)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_borders(cell)
    doc.add_paragraph().paragraph_format.space_after = Pt(3)

def build_markdown(src, out_path, blinded=True):
    doc = Document()
    configure(doc)
    lines = Path(src).read_text().splitlines()
    i = 0
    while i < len(lines):
        line = lines[i]
        if not line.strip():
            i += 1
            continue
        if line.startswith("|"):
            block = []
            while i < len(lines) and lines[i].startswith("|"):
                cells = [c.strip() for c in lines[i].strip().strip("|").split("|")]
                if not all(set(c) <= set("-: ") for c in cells):
                    block.append(cells)
                i += 1
            if block:
                add_table(doc, block)
            continue
        heading = re.match(r"^(#{1,6})\s+(.*)$", line)
        if heading:
            level = min(len(heading.group(1)), 3)
            text = re.sub(r"[`*_]", "", heading.group(2))
            p = doc.add_paragraph(style="Title" if level == 1 and not any(p.style.name == "Title" for p in doc.paragraphs) else f"Heading {level}")
            add_inline(p, text)
            i += 1
            continue
        if re.match(r"^\d+\.\s+", line):
            p = doc.add_paragraph(style="List Number")
            p.paragraph_format.line_spacing = WD_LINE_SPACING.DOUBLE
            add_inline(p, re.sub(r"^\d+\.\s+", "", line))
            i += 1
            continue
        if re.match(r"^-\s+", line):
            p = doc.add_paragraph(style="List Bullet")
            p.paragraph_format.line_spacing = WD_LINE_SPACING.DOUBLE
            add_inline(p, re.sub(r"^-\s+", "", line))
            i += 1
            continue
        p = doc.add_paragraph()
        add_inline(p, line)
        i += 1
    doc.save(out_path)

def build_metadata(src, out_path, title_page=False):
    doc = Document()
    configure(doc)
    if title_page:
        normal = doc.styles["Normal"]
        normal.font.size = Pt(11)
        normal.paragraph_format.line_spacing = 1.0
        normal.paragraph_format.space_after = Pt(2)
    lines = Path(src).read_text().splitlines()
    for line in lines:
        if not line.strip():
            continue
        if title_page and (line.startswith("Use this as") or line.startswith("identifying information")):
            continue
        heading = re.match(r"^(#{1,6})\s+(.*)$", line)
        if heading:
            level = min(len(heading.group(1)), 3)
            text = re.sub(r"[`*_]", "", heading.group(2))
            if title_page and text == "JPS title-page template":
                text = "Title Page"
            p = doc.add_paragraph(style="Title" if level == 1 else f"Heading {level}")
            add_inline(p, text)
        else:
            p = doc.add_paragraph()
            add_inline(p, line)
    doc.save(out_path)

build_markdown(ROOT / "manuscript" / "article.md", OUT / "absinthe-jps-blinded-manuscript.docx")
build_metadata(ROOT / "manuscript" / "jps-title-page-template.md", OUT / "absinthe-jps-title-page.docx", True)
build_metadata(ROOT / "manuscript" / "jps-cover-letter.md", OUT / "absinthe-jps-cover-letter.docx")
print("created", *(p.name for p in OUT.glob("*.docx")))
